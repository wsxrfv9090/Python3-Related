from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, fill='FFA500'):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)

# Map day numbers to circled Unicode characters.
# 1-20: ① to ⑳ (U+2460 to U+2473)
# 21-31 use: 21=㉑(U+3251), 22=㉒(U+3252) ... 31=㉛(U+325B)
circled_nums = {
    1:"①",2:"②",3:"③",4:"④",5:"⑤",6:"⑥",7:"⑦",8:"⑧",9:"⑨",10:"⑩",
    11:"⑪",12:"⑫",13:"⑬",14:"⑭",15:"⑮",16:"⑯",17:"⑰",18:"⑱",19:"⑲",20:"⑳",
    21:"㉑",22:"㉒",23:"㉓",24:"㉔",25:"㉕",26:"㉖",27:"㉗",28:"㉘",29:"㉙",30:"㉚",31:"㉛"
}

def set_run_font(run, size=14, bold=False, color=(255, 255, 255)):
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(color[0], color[1], color[2])

doc = Document()

# Set page orientation to landscape
section = doc.sections[-1]
section.orientation = WD_ORIENT.LANDSCAPE
new_width, new_height = section.page_height, section.page_width
section.page_width = new_width
section.page_height = new_height
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)

# Create main table
main_table = doc.add_table(rows=1, cols=2)
main_table.autofit = False
main_table.columns[0].width = Inches(5)
main_table.columns[1].width = Inches(5)

# Apply shading to both cells
for cell in main_table.row_cells(0):
    set_cell_shading(cell, 'FFA500')  # Orange background

left_cell = main_table.cell(0,0)
right_cell = main_table.cell(0,1)

# Left side content
p = left_cell.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
title_run = p.add_run('2020 “三伏” 日历')
set_run_font(title_run, size=24, bold=True, color=(255,255,255))

p = left_cell.add_paragraph()
exp_run = p.add_run('“三伏”是初伏、中伏和末伏的总称。三伏恰在小暑和大暑之间，是一年中最热的时候。')
set_run_font(exp_run, size=12, bold=False, color=(255,255,255))

# July Calendar Title
july_para = left_cell.add_paragraph()
july_title_run = july_para.add_run('2020年7月日历（头伏、中伏）')
set_run_font(july_title_run, size=14, bold=True, color=(255,255,255))

# Colors for different伏
head_color = (255, 175, 100)   # 头伏
middle_color = (255, 100, 0)   # 中伏
end_color = (200, 0, 0)        # 末伏

# July data (based on image)
# Adjust as per the original pattern (example from previous code).
weeks_july = [
    ([6,7,8,9,10,11], head_color),
    ([12,13,14,15,16,17], head_color),
    ([18,19,20,21,22,23], middle_color),
    ([24,25,26,27,28,29], middle_color),
    ([30,31], middle_color)
]

for week, color_set in weeks_july:
    week_p = left_cell.add_paragraph()
    for d in week:
        day_run = week_p.add_run(circled_nums[d] + ' ')
        set_run_font(day_run, size=12, bold=True, color=color_set)

# August Calendar Title
aug_para = left_cell.add_paragraph()
aug_title_run = aug_para.add_run('2020年8月日历（末伏）')
set_run_font(aug_title_run, size=14, bold=True, color=(255,255,255))

weeks_aug = [
    ([1,2,3,4,5,6], middle_color),
    ([7,8,9,10,11,12], middle_color),
    ([13,14,15,16,17,18], end_color),
    ([19,20,21,22,23,24], end_color)
]

for week, color_set in weeks_aug:
    week_p = left_cell.add_paragraph()
    for d in week:
        day_run = week_p.add_run(circled_nums[d] + ' ')
        set_run_font(day_run, size=12, bold=True, color=color_set)

# Right side content
# The user provided text:
right_text = (
    "2020年7月16日至7月25日为初伏，7月26日至8月14日为中伏，8月15日至8月24日为末伏。今年三伏长达40天。\n"
    "三伏的日期是按节气的日期和干支的日期相配合来决定的。按农历的规定，夏至以后的第三个庚日（干支纪日法中带“庚”的日子称为庚日）为初伏（也叫头伏）；"
    "第四个庚日为中伏（也叫二伏）；立秋后的第一个庚日为末伏（也叫三伏）。\n"
    "通过如下的函数建立的极端气象模型，可以在排除偶然性因素的前提下，分析出不同位置极端数值天气出现的趋势。\n"
)

p_right = right_cell.add_paragraph()
run_r = p_right.add_run(right_text)
set_run_font(run_r, size=12, bold=False, color=(255,255,255))

# Insert the formula for 广义极值分布函数
p_formula_title = right_cell.add_paragraph()
ft_run = p_formula_title.add_run('广义极值分布函数\n')
set_run_font(ft_run, size=14, bold=True, color=(255,255,255))

desc_run = p_formula_title.add_run(
    '通过如下的函数表述出的极端气象模型，可以在排除偶然性因素的前提下，分析出不同位置极端数值天气出现的趋势：\n'
)
set_run_font(desc_run, size=12, bold=False, color=(255,255,255))

# Formula lines
# We'll write them as multiple runs for clarity
formula_p = right_cell.add_paragraph()
formula_lines = [
    "          { exp{-[1 - k(x - ε)/a]^(1/k)},   k<0, x ≥ ε + a/k",
    "F(x) = { exp{-exp[-(x - ε)/δ]}          , k=0",
    "          { exp{-[1 - k(x - ε)/a]^(1/k)},   k>0, x ≤ ε + a/k"
]

for line in formula_lines:
    fr = formula_p.add_run(line + "\n")
    set_run_font(fr, size=11, bold=False, color=(255,255,255))

doc.save('sanfu_calendar.docx')
